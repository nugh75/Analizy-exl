import pandas as pd
import logging
import os
from langchain_openai import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from docx import Document
from dotenv import load_dotenv

# Carica le variabili di ambiente dal file .env
load_dotenv()

# Estrai la chiave API di OpenAI dall'ambiente
openai_api_key = os.getenv('OPENAI_API_KEY')

if not openai_api_key:
    raise ValueError("La chiave API di OpenAI non è stata trovata. Assicurati che sia presente nel file .env.")

# Configura il modello OpenAI tramite LangChain usando ChatOpenAI con il modello gpt-4o-mini
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.7, api_key=openai_api_key)

# Configurazione del logging per mostrare anche nel terminale
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("operazioni_log.log"),
        logging.StreamHandler()  # Aggiungi il log sul terminale
    ]
)

# Log dell'inizio del processo
logging.info("Inizio del processo di analisi dei debbolezze")

# Fase 1: Estrazione dei debbolezze da ciascun commento
prompt_template_1 = ChatPromptTemplate.from_template(
    """Analizza il seguente commento che descrive i debbolezze del progetto "scuole aperte". Riassumi i punti principali emersi:
    {commento}
    """
)

# Definizione del prompt per la Fase 2 (sintesi complessiva)
prompt_template_2 = ChatPromptTemplate.from_template(
    """Ecco un elenco di debbolezze estratti da vari commenti di insegnanti del progetto "scuole aperte":
    {punti_di_DEBOLEZZA}
    Analizza questi debbolezze, raggruppa quelli simili e fai una sintesi dei temi principali emersi.
    """
)

# Definizione del prompt per la Fase 3 (selezione esempi per ciascun DEBOLEZZA)
prompt_template_3 = ChatPromptTemplate.from_template(
    """Dato il seguente elenco di debbolezze:
    {punti_di_DEBOLEZZA}
    
    Seleziona per ciascun DEBOLEZZA 10 esempi di commenti dal file originale. Per ogni esempio, includi l'ID del commento e il testo del commento.
    """
)

# Leggi il file Excel
file_path = "/Users/desi76/Library/CloudStorage/OneDrive-UniversitadegliStudiRomaTre/Scuole aperte file non condivisi/debolezze.xlsx"  # Cambia con il percorso corretto
try:
    df = pd.read_excel(file_path)
    logging.info(f"File Excel '{file_path}' caricato con successo.")
except Exception as e:
    logging.error(f"Errore durante il caricamento del file Excel: {e}")
    raise

# Lista per salvare i risultati dell'analisi (Fase 1)
report_debbolezze = []

# Fase 1: Itera su ogni riga della colonna 'DEBOLEZZA' e estrai i debbolezze
for index, row in df.iterrows():
    id_commento = row['id']  # ID del commento
    commento = row['DEBOLEZZA']  # Testo del commento da analizzare
    
    # Verifica se la cella del commento è vuota
    if pd.isna(commento):
        logging.info(f"Cella vuota per l'ID {id_commento}, saltato.")
        report_debbolezze.append("Cella vuota")  # Mantieni una voce per le celle vuote
        continue
    
    try:
        # Log dell'inizio dell'analisi del commento
        logging.info(f"Analisi del commento con ID {id_commento}: {commento}")
        
        # Crea il prompt per la prima fase
        prompt = prompt_template_1.format(commento=commento)

        # Crea la struttura del messaggio richiesta da OpenAI
        messages = [{"role": "system", "content": "Sei un assistente che analizza i debbolezze del progetto 'scuole aperte'."},
                    {"role": "user", "content": prompt}]
        
        # Invoca l'analisi tramite LangChain/OpenAI per estrarre i debbolezze
        result = llm.invoke(messages)  # Usa il metodo `invoke` al posto di `__call__`
        
        # Aggiungi il risultato alla lista dei debbolezze
        report_debbolezze.append(result.content)  # Usa `content` per accedere al testo
        
        # Log del successo dell'analisi del commento
        logging.info(f"Risultato dell'analisi per l'ID {id_commento}: {result.content}")
    except Exception as e:
        # Log di eventuali errori durante l'analisi
        logging.error(f"Errore durante l'analisi del commento con ID {id_commento}: {e}")
        report_debbolezze.append(f"Errore nell'analisi: {e}")

# Crea un nuovo DataFrame per il report dei debbolezze estratti
df_report = pd.DataFrame({
    'ID': df['id'],  # Colonna con l'ID
    'Commento Originale': df['DEBOLEZZA'],  # Colonna con i commenti originali
    'debbolezze': report_debbolezze  # Lista dei debbolezze
})

# Salva il report in un nuovo file Excel (Fase 1)
try:
    df_report.to_excel('report_debbolezze.xlsx', index=False)
    logging.info("Report della Fase 1 salvato con successo in 'report_debbolezze.xlsx'.")
except Exception as e:
    logging.error(f"Errore durante il salvataggio del report della Fase 1: {e}")
    raise

# Fase 2: Analisi dei debbolezze complessivi
punti_di_DEBOLEZZA_complessivi = "\n".join([p for p in report_debbolezze if p != "Cella vuota"])

try:
    logging.info("Inizio dell'analisi complessiva dei debbolezze.")
    
    # Crea il prompt per la seconda fase
    prompt_completo = prompt_template_2.format(punti_di_DEBOLEZZA=punti_di_DEBOLEZZA_complessivi)

    # Crea la struttura del messaggio per la seconda fase
    messages_completi = [{"role": "system", "content": "Sei un assistente che analizza i debbolezze del progetto 'scuole aperte'."},
                        {"role": "user", "content": prompt_completo}]
    
    # Invoca l'analisi di LangChain/OpenAI per raggruppare e sintetizzare i debbolezze
    sintesi_debbolezze = llm.invoke(messages_completi)  # Usa `invoke`
    
    # Log del successo dell'analisi complessiva
    logging.info("Analisi complessiva dei debbolezze completata con successo.")
except Exception as e:
    logging.error(f"Errore durante l'analisi complessiva dei debbolezze: {e}")
    raise

# Fase 3: Selezione di 5 esempi per ogni DEBOLEZZA (ID e commento)
try:
    logging.info("Inizio della selezione di 5 esempi per ciascun DEBOLEZZA.")
    
    # Crea il prompt per la terza fase
    prompt_selezione = prompt_template_3.format(punti_di_DEBOLEZZA=punti_di_DEBOLEZZA_complessivi)

    # Crea la struttura del messaggio per la terza fase
    messages_selezione = [{"role": "system", "content": "Sei un assistente che seleziona esempi di commenti per ciascun DEBOLEZZA."},
                        {"role": "user", "content": prompt_selezione}]
    
    # Invoca l'analisi di LangChain/OpenAI per selezionare 5 esempi per ciascun DEBOLEZZA
    selezione_esempi = llm.invoke(messages_selezione)  # Usa `invoke`
    
    # Log del successo della selezione degli esempi
    logging.info("Selezione degli esempi completata con successo.")
except Exception as e:
    logging.error(f"Errore durante la selezione degli esempi per i debbolezze: {e}")
    raise

# Rimuovi eventuali segni markdown dai risultati
def remove_markdown(text):
    return text.replace("**", "").replace("##", "")

# Salva la sintesi e gli esempi in un file di testo
try:
    with open('sintesi_debbolezze_e_esempi.txt', 'w') as file:
        # Scrivi la sintesi dei debbolezze
        file.write(remove_markdown(sintesi_debbolezze.content))
        
        # Scrivi i 5 esempi per ciascun DEBOLEZZA
        file.write("\n\nEsempi di Commenti per ciascun DEBOLEZZA:\n")
        file.write(remove_markdown(selezione_esempi.content))
    
    logging.info("Sintesi ed esempi salvati in 'sintesi_debbolezze_e_esempi.txt'.")
except Exception as e:
    logging.error(f"Errore durante il salvataggio della sintesi e degli esempi: {e}")
    raise

# Salva la sintesi e gli esempi in un file docx
try:
    doc = Document()
    
    # Aggiungi la sintesi dei debbolezze al documento Word
    doc.add_heading('Sintesi dei debbolezze', level=1)
    doc.add_paragraph(remove_markdown(sintesi_debbolezze.content))
    
    # Aggiungi gli esempi al documento Word
    doc.add_heading('Esempi di Commenti per ciascun DEBOLEZZA', level=1)
    doc.add_paragraph(remove_markdown(selezione_esempi.content))
    
    # Salva il file Word
    doc.save('sintesi_debbolezze_e_esempi.docx')
    logging.info("Sintesi ed esempi salvati in 'sintesi_debbolezze_e_esempi.docx'.")
except Exception as e:
    logging.error(f"Errore durante il salvataggio del file Word: {e}")
    raise

logging.info("Processo completato con successo.")
print("Analisi completata. Report della Fase 1 salvato in 'report_debbolezze.xlsx'.")
print("Sintesi ed esempi salvati in 'sintesi_debbolezze_e_esempi.txt' e 'sintesi_debbolezze_e_esempi.docx'.")
